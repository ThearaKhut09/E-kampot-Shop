#!/usr/bin/env python3
"""
Generate a single-page DOCX checklist from docs/TEST_CHECKLIST.md
Requires: python-docx

Usage:
    pip install python-docx
    python docs/generate_checklist_docx.py

Output: docs/TEST_CHECKLIST.docx
"""
from docx import Document
from docx.shared import Pt
import os

ROOT = os.path.dirname(os.path.realpath(__file__))
INPUT = os.path.join(ROOT, 'TEST_CHECKLIST.md')
OUTPUT = os.path.join(ROOT, 'TEST_CHECKLIST.docx')

if not os.path.exists(INPUT):
    print(f"Input file not found: {INPUT}")
    raise SystemExit(1)

with open(INPUT, 'r', encoding='utf-8') as f:
    lines = [l.rstrip() for l in f.readlines()]

# Create docx
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = lines[0] if lines else 'Test Checklist'
doc.add_heading(title.replace('# ', ''), level=1)

# Add remaining lines as bullets
for line in lines[2:]:
    if not line.strip():
        continue
    if line.lstrip().startswith('- '):
        text = line.lstrip()[2:].strip()
        doc.add_paragraph(text, style='List Bullet')
    else:
        doc.add_paragraph(line)

# Save
if os.path.exists(OUTPUT):
    os.remove(OUTPUT)

doc.save(OUTPUT)
print(f'Generated: {OUTPUT}')
